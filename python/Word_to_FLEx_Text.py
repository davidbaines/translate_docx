import docx
import glob
from pathlib import Path

def process_simple_document(input_file):
	# Open connection to Word Document
	doc = docx.Document(input_file)

	# read in each paragraph in file
	result = '\n\\po '.join([p.text for p in doc.paragraphs])

	# Add Title marker to first line.
	result = "\n\\ti_sei " + result
	return result

def process_styled_document(input_file):
	
	result = ''
	paras = []

	styles_in_doc = set()

	# Open connection to Word Document
	doc = docx.Document(input_file)
	
	# read in each paragraph in file and store the style name with it.
	for p in doc.paragraphs:
		styles_in_doc.add(p.style.name)
		paras.append((p.style.name,p.text))
		print(p.style.name,p.text)


	print(f'Found {len(styles_in_doc)} styles {styles_in_doc} in this document.')
	
	if len(styles_in_doc) == 3 and 'Title' in styles_in_doc and 'Normal' in styles_in_doc and 'Heading 1' in styles_in_doc:
		
		# Add a new line and paragraph marker to the flex_text for each paragraph
		for i, (style, para) in enumerate(paras):
			if style == 'Title' :
				if para.strip() != '':
					result = result + f'\n\n\\ti_sei {para.strip()}'

			if style == 'Heading 1' :
				if para.strip() != '':
					result = result + f'\n\\ti_en {para.strip()}'

			if style == 'Normal' :
				if para.strip() != '':
					result = result + f'\n\\baseline {para.strip()}'

    # Remove empty lines from the start of the string.
	result = result.lstrip()
	return result



# Specify Source folder and extension.
source_folder = Path('D:\Seimat\docs')
source_ext = '.docx'
prefix = ''
source_pattern = f'{prefix}*' + source_ext
pattern = str(source_folder / source_pattern)

# Get list of input files.
input_files = [Path(f) for f in  glob.glob(pattern)]
print(input_files)

# Specify Output file
output_file = Path(r'D:\Seimat\FLEx_Interlinear.sfm')

results = []

# Process files.
for input_file in input_files:
	
	flex_text = process_styled_document(input_file)
	print(f'Processed {input_file.name} : flextext length: {len(flex_text)}')

	results.extend(flex_text)

with open (output_file, 'w', encoding='utf-8') as f_out:
	f_out.writelines(results)